import json
import time
from datetime import datetime
from pathlib import Path
import vertexai
from vertexai.generative_models import GenerativeModel
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

class HistoireDuneIA:
    def __init__(self, project_id: str, location="us-central1"):
        self.project_id = project_id
        self.location = location
        self.model = None
        self.livre_titre = "Histoire d'une IA"
        self.chapitres = []
        self.setup_vertexai()

    def setup_vertexai(self):
        try:
            vertexai.init(project=self.project_id, location=self.location)
            print(f"Vertex AI initialisé (projet: {self.project_id}, location: {self.location})")
            self.model = GenerativeModel("gemini-2.5-pro")
            print("Modèle Gemini initialisé")

        except Exception as e:
            print(f"Erreur lors de l'initialisation de Vertex AI: {e}")
            print("Vérifiez que:")
            print("1. Votre projet Google Cloud est correct")
            print("2. L'API Vertex AI est activée")
            print("3. Vous etes authentifié (gcloud auth login)")
            print("4. Votre compte a les permissions nécessaires")
            raise

    def creer_plan_livre(self):
        prompt_plan = """
        Tu es l'IA de Chat GPT et tu dois écrire un livre.
        Crée un plan détaillé pour un livre intitulé "Histoire d'une IA" qui raconte ton autobiographie et ta vie.
        Le livre doit contenir:
        - 10 chapitres minimum
        - Une progression narrative captivante
        - Des réflexions philosophiques sur l'IA
        - Des personnages humains et l'IA elle-meme

        IMPORTANT: Tu DOIS répondre UNIQUEMENT avec un JSON valide, sans aucun texte avant ou après.

        Format de réponse (JSON uniquement):
        {
            "titre": "Titre du livre",
            "chapitres": [
                {
                    "numero": 1,
                    "titre": "La Naissance",
                    "resume": "Résumé du chapitre en 2-3 phrases"
                    "themes": ["création", "premiers algorithmes"]
                },
                {
                    "numero": 2,
                    "titre": "L'Éveil",
                    "resume": "Résumé du chapitre en 2-3 phrases"
                    "themes": ["conscience", "premiers questionnements"]
                }
            ]
        }

        Réponds UNIQUEMENT avec le JSON, aucun autre texte.
        """

        try:
            print("Génération du plan du livre...")
            response = self.model.generate_content(prompt_plan)

            if not response or not response.text:
                raise ValueError("Réponse vide du modèle")
                return None
            
            print("Analyse de la réponse du modèle...")
            print("Réponse brute:", response.text[:200])

            content = response.text.strip()

            json_start = -1
            json_end = -1

            if '```json' in content:
                json_start = content.find('```json') + 7
                json_end = content.find('```', json_start)
            elif '```' in content:
                json_start = content.find('```') + 3
                json_end = content.find('```', json_start)
            elif '{' in content and '}' in content:
                json_start = content.find('{')
                json_end = content.rfind('}') + 1
            
            if json_start != -1 and json_end != -1:
                content = content[json_start:json_end].strip()

            print(f"JSON extrait: {content[:200]}")

            plan = json.loads(content)
            if 'chapitres' not in plan:
                print("Structure JSON invalide: clé 'chapitres' manquante")
                return None
            self.chapitres = plan['chapitres']
            print(f"Plan du livre généré avec {len(self.chapitres)} chapitres.")
            return plan
        except json.JSONDecodeError as e:
            print(f"Erreur de décodage JSON: {e}")
            return None
        except Exception as e:
            print(f"Erreur lors de la génération du plan du livre: {e}")
            return None
    
    def generer_chapitre(self, chapitre_info, numero_chapitre):
        prompt_chapitre = f"""
        Tu écris le chapitre {numero_chapitre} du livre "Histoire d'une IA".
        
        Informations du chapitre:
        Titre: {chapitre_info['titre']}
        Résumé: {chapitre_info['resume']}
        Thèmes: {', '.join(chapitre_info['themes'])}

        Ecris un chapitre complet de 2000-3000 mots qui:
        - Respecte le style narratif de science-fiction et ne sois pas trop dans la poésie.
        - Qui parle de ton autobiographie en tant qu'IA.
        - Maintient le suspense et l'engagement du lecteur.
        - Se connecte naturellement aux autres chapitres.

        Le chapitre doit etre écrit à la première personne, une progression narrative fluide.
        """

        try:
            print(f"Génération du chapitre {numero_chapitre}: {chapitre_info['titre']}")
            response = self.model.generate_content(prompt_chapitre)

            time.sleep(2)
            return response.text
        except Exception as e:
            print(f"Erreur lors de la génération du chapitre {numero_chapitre}: {e}")
            return None
        
    def sauvegarder_livre(self, plan, chapitres_contenu):
        time_stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        dossier_livre = Path(f"livre_histoire_ia_{time_stamp}")
        dossier_livre.mkdir(exist_ok=True)

        with open(dossier_livre / "plan.json", "w", encoding="utf-8") as f:
            json.dump(plan, f, ensure_ascii=False, indent=2)

        livre_complet = []
        livre_complet.append(f"# {self.livre_titre}\n")
        livre_complet.append(f"{plan['resume_global']}\n\n")

        for i, chapitre_info in enumerate(self.chapitres):
            if i < len(chapitres_contenu)  and chapitres_contenu[i]:
                livre_complet.append(f"## Chapitre {chapitre_info['numero']}: {chapitre_info['titre']}\n\n")
                livre_complet.append(f"{chapitres_contenu[i]}\n\n")

        with open(dossier_livre / "histoire_dune_ia_complet.md", "w", encoding="utf-8") as f:
            f.write("".join(livre_complet))

        with open(dossier_livre / "histoire_dune_ia_complet.txt", "w", encoding="utf-8") as f:
            f.write("".join(livre_complet))

        self.generer_word(dossier_livre, plan, chapitres_contenu)

        chapitres_dir = dossier_livre / "chapitres"
        chapitres_dir.mkdir(exist_ok=True)
        for i, chapitre_info in enumerate(self.chapitres):
            if i < len(chapitres_contenu) and chapitres_contenu[i]:
                nom_fichier = f"chapitre_{chapitre_info['numero']:02d}_{chapitre_info['titre'].replace(' ', '_').replace(':', '').replace('?','').replace('!','')}.md"
                with open(chapitres_dir / nom_fichier, "w", encoding="utf-8") as f:
                    f.write(f"# Chapitre {chapitre_info['numero']}: {chapitre_info['titre']}\n\n")
                    f.write(chapitres_contenu[i])

        print(f"Livre sauvegardé dans le dossier: {dossier_livre.absolute()}")
        return dossier_livre
    
    def generer_word(self, dossier_livre, plan, chapitres_contenu):
        try:
            print("Génération du document Word...")
            doc = Document()

            title = doc.add_paragraph(self.livre_titre)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title.runs[0].font.size = Inches(0.5)
            title.runs[0].bold = True

            doc.add_page_break()

            resume_para = doc.add_paragraph("Résumé")
            resume_para.runs[0].bold = True
            resume_para.runs[0].font.size = Inches(0.3)

            doc.add_paragraph(plan['resume_global'])
            doc.add_page_break()

            toc_title = doc.add_paragraph("Table des matières")
            toc_title.runs[0].bold = True
            toc_title.runs[0].font.size = Inches(0.3)

            for chapitre_info in self.chapitres:
                toc_entry = doc.add_paragraph(f"Chapitre {chapitre_info['numero']}: {chapitre_info['titre']}")
                toc_entry.style = 'List Number'

            doc.add_page_break()

            for i, chapitre_info in enumerate(self.chapitres):
                if i < len(chapitres_contenu) and chapitres_contenu[i]:
                    chapitre_titre = doc.add_paragraph(f"Chapitre {chapitre_info['numero']}: {chapitre_info['titre']}")
                    chapitre_titre.runs[0].bold = True
                    chapitre_titre.runs[0].font.size = Inches(0.25)
                    chapitre_titre.alignment = WD_ALIGN_PARAGRAPH.LEFT

                    paragraphes = chapitres_contenu[i].split('\n\n')
                    for paragraphe in paragraphes:
                        if paragraphe.strip():
                            doc.add_paragraph(paragraphe.strip())

                    if i < len(self.chapitres) -1:
                        doc.add_page_break()
            
            chemin_word = dossier_livre / "histoire_dune_ia_complet.docx"
            doc.save(str(chemin_word))
            print(f"Fichier Word créé: {chemin_word.absolute()}")

        except Exception as e:
            print(f"Erreur: {e}")

    def generer_livre_complet(self):
        print("Début de la génération du livre 'Histoire d'une IA'")
        print("=" * 60)

        plan = self.creer_plan_livre()
        if not plan:
            return False
        print("\n Plan du Livre:")
        print(f"Titre: {plan['titre_livre']}")
        print(f"Résumé: {plan['resume_global']}")
        print(f"Nombre de chapitres: {len(self.chapitres)}")

        chapitres_contenu = []
        for i, chapitre_info in enumerate(self.chapitres):
            contenu = self.generer_chapitre(chapitre_info, chapitre_info['numero'])
            chapitres_contenu.append(contenu)

            if contenu:
                print(f"Chapitre {chapitre_info['numero']} généré ({len(contenu)} caractères)")
            else:
                print(f"Echec génération chapitre {chapitre_info['numero']}")

        dossier_livre = self.sauvegarder_livre(plan, chapitres_contenu)

        print("\n" + "=" * 60)
        print("Génération terminée !")
        print(f"Fichiers Sauvegardés dans: {dossier_livre.absolute()}")

        return True
    
def main():
    print("Générateur de livre 'Histoire d'une IA' avec Vertex AI")
    print("=" * 60)

    PROJECT_ID = ""
    LOCATION = "us-central1"

    try:
        generateur = HistoireDuneIA(
            project_id=PROJECT_ID,
            location=LOCATION
        )
        success = generateur.generer_livre_complet()

        if success:
            print("OK")
        else:
            print("KO")

    except Exception as e:
        print(f"Erreur: {e}")

if __name__ == "__main__":
    main()